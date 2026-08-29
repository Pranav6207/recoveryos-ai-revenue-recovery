from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "AI_Revenue_Recovery_SRS.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "64748B"
LIGHT = "F2F4F7"
CALLOUT = "FFF8E8"
GOLD = "B26B00"
GREEN = "0F766E"
RED = "9B1C1C"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_dxa / 1440)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(TABLE_WIDTH))
    table_width.set(qn("w:type"), "dxa")
    table_indent = tbl_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        tbl_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT))
    table_indent.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        grid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def set_keep_with_next(paragraph):
    paragraph_pr = paragraph._p.get_or_add_pPr()
    element = OxmlElement("w:keepNext")
    paragraph_pr.append(element)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        lead = paragraph.add_run(bold_prefix)
        set_font(lead, size=11, color=INK, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_font(rest, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11, color=INK)
    return paragraph


def add_h1(doc, text):
    paragraph = doc.add_paragraph(style="Heading 1")
    run = paragraph.add_run(text)
    set_font(run, size=16, color=BLUE, bold=True)
    set_keep_with_next(paragraph)
    return paragraph


def add_h2(doc, text):
    paragraph = doc.add_paragraph(style="Heading 2")
    run = paragraph.add_run(text)
    set_font(run, size=13, color=BLUE, bold=True)
    set_keep_with_next(paragraph)
    return paragraph


def add_callout(doc, title, text, accent=GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_shading(header_cells[index], LIGHT)
        p = header_cells[index].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, size=9.5, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_font(run, size=9.5, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("RECOVERYOS  |  SYSTEM REQUIREMENTS SPECIFICATION")
    set_font(header_run, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(26)
    kicker.paragraph_format.space_after = Pt(4)
    r = kicker.add_run("RAZORPAY AI BUILDATHON · TRACK 03")
    set_font(r, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("RecoveryOS")
    set_font(r, size=29, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("System Requirements Specification")
    set_font(r, size=15, color=DARK_BLUE, bold=True)

    add_table(doc, ["Document control", "Value"], [
        ("Version", "1.0 · Hackathon Selection Release"),
        ("Prepared for", "Razorpay AI Buildathon judges and the RecoveryOS build team"),
        ("Release goal", "A zero-configuration, judge-testable AI revenue recovery demonstration"),
        ("Date", date(2026, 8, 29).strftime("%d %B %Y")),
    ], [2700, 6660])

    add_callout(doc, "Purpose", "RecoveryOS proves a complete recovery loop: detect revenue at risk, diagnose the cause, select a compliant intervention, execute a bounded workflow, and report confirmed recovered value with evidence.")
    add_body(doc, "This specification defines the buildathon release. It intentionally excludes live collections operations, real customer data, and autonomous movement of money.")
    doc.add_page_break()

    add_h1(doc, "1. Product overview")
    add_body(doc, "RecoveryOS is a responsive, installable web command center for revenue-risk recovery. It turns payment and receivable signals into reviewable cases, combines an AI recommendation with deterministic policy checks, then runs a safe simulator or optional test-mode adapter.")
    add_h2(doc, "1.1 Product objectives")
    add_table(doc, ["Objective", "Success condition"], [
        ("Show all recovery loops", "A judge can inspect and execute a scenario for every challenge direction."),
        ("Prove real value", "Batch metrics report at-risk value, confirmed recovered value, recovery rate, and excluded cases."),
        ("Show AI safely", "Recommendations are explainable; policy code controls execution and stop conditions."),
        ("Remove setup friction", "The hosted and local demos run without Razorpay, Resend, Supabase, or LLM credentials."),
    ], [2700, 6660])
    add_h2(doc, "1.2 Scope")
    add_table(doc, ["Included", "Excluded"], [
        ("Synthetic 120-case batch and isolated demo run", "Real customer PII, live collections, billing, and merchant onboarding"),
        ("Seven recovery playbooks and a browser voice preview", "Actual calls, SMS, mandate debits, or autonomous money movement"),
        ("Audit events, policy blocks, recovery attribution, analytics", "Enterprise SLA, complex roles, multi-tenant administration, and paid scheduler dependencies"),
        ("Optional Razorpay Test Mode, Supabase, and Resend adapters", "Production payment credentials in the public demo"),
    ], [4680, 4680])

    add_h1(doc, "2. Users and workflows")
    add_table(doc, ["Actor", "Primary need", "Permitted interaction"], [
        ("Judge", "Evaluate product depth in minutes", "Start a private demo run, inspect evidence, and execute safe demo actions"),
        ("Demo operator", "Present recovery value", "Filter cases, approve bounded actions, preview voice script, and advance demo time"),
        ("Recovery policy engine", "Protect customers and revenue operations", "Allow, defer, or block each proposed action with a visible reason"),
    ], [1800, 3960, 3600])

    add_h2(doc, "2.1 Required recovery playbooks")
    add_table(doc, ["ID", "Trigger and diagnosis", "Bounded recovery action"], [
        ("FR-01 Payment degradation", "Repeated payment failure or declining authorization pattern", "Create a fresh verified payment-link recovery path"),
        ("FR-02 Checkout drop-off", "Checkout remains incomplete after a configured timeout", "Send one consent-approved checkout recovery message"),
        ("FR-03 Subscription failure", "Subscription enters pending or halted state", "Respect provider retry state; request a payment-method update when appropriate"),
        ("FR-04 B2B receivable", "Invoice remains overdue without a dispute or settlement", "Prioritize the invoice and issue/reuse a traceable payment link"),
        ("FR-05 Mandate retry", "Mandate charge fails and a provider-approved window exists", "Recommend the next safe retry window; never initiate a duplicate debit"),
        ("FR-06 Hinglish voice", "Consent exists and a high-value case needs assisted recovery", "Show reviewed Hinglish script, browser voice preview, and simulated outcome"),
        ("FR-07 Promise-to-pay", "Customer commits to a payment amount/date", "Record the promise, suppress new chasing, and reopen only after expiry unpaid"),
    ], [1800, 3780, 3780])

    add_h1(doc, "3. Functional requirements")
    add_table(doc, ["ID", "Requirement"], [
        ("FR-08", "The system shall generate a 120-case deterministic synthetic batch spanning all seven playbooks."),
        ("FR-09", "The system shall create an isolated demo run so a judge can interact without corrupting the baseline experience."),
        ("FR-10", "The system shall display each case's amount at risk, risk score, confidence, root cause, suggested channel, and state."),
        ("FR-11", "The system shall issue a structured recommendation containing the action, rationale, approved channel, communication draft, and expected outcome."),
        ("FR-12", "The system shall write an audit event for each signal, AI recommendation, policy decision, action, outcome, and recovery attribution."),
        ("FR-13", "The system shall count recovered amount only after a confirmed simulated or provider payment event and shall attribute that payment only once."),
        ("FR-14", "The system shall provide batch metrics for at-risk amount, recovered amount, recovery rate, policy-protected cases, and executed actions."),
        ("FR-15", "The system shall provide an explicit advance-demo-time action instead of depending on a background scheduler."),
    ], [1800, 7560])

    add_h1(doc, "4. Policy, AI, and audit controls")
    add_body(doc, "The intelligence layer may explain a root cause and select from approved playbooks, but it may not directly debit an account, alter a mandate, or bypass a policy decision.")
    add_table(doc, ["Control", "Rule", "System behavior"], [
        ("Consent", "No valid permission for the selected outreach channel", "Block the action and preserve the case for review"),
        ("Opt-out", "Customer preference prohibits automated outreach", "Stop all automated contact attempts"),
        ("Contact cap", "Three automated attempts within seven days", "Defer the case or route it to a human reviewer"),
        ("Quiet hours", "Configured customer quiet period is active", "Queue the action for the next permitted window"),
        ("Settlement", "A recovered payment is confirmed", "Close the case and stop all follow-up"),
        ("Promise-to-pay", "A valid promise date has not expired", "Pause reminders until payment or expiry"),
        ("Provider retry", "Subscription/mandate retry is still controlled by provider", "Do not create a competing charge or duplicate debit"),
    ], [1800, 3600, 3960])
    add_callout(doc, "Audit principle", "Every approval, refusal, simulator outcome, and recovery attribution is visible in the selected case timeline. This lets judges inspect why the agent acted, not only what it acted on.", GREEN)

    add_h1(doc, "5. Solution architecture")
    add_body(doc, "Event source → recovery case → structured AI recommendation → policy gate → safe adapter → reconciliation → audit and batch metrics.")
    add_table(doc, ["Component", "Responsibility", "Release implementation"], [
        ("Next.js PWA", "Responsive command center and server route surface", "App Router, TypeScript, Tailwind CSS, installable manifest"),
        ("Recovery engine", "Risk, recommendation, policy, and outcome logic", "Deterministic TypeScript functions covered by Vitest"),
        ("Simulator adapters", "Reliable zero-config outcomes", "Payment, email, voice, checkout, mandate, and promise behavior"),
        ("Razorpay adapter", "Optional Test Mode webhook/payment-link interface", "HMAC-validating webhook route and payment-link helper"),
        ("Supabase adapter", "Optional persistent audit storage", "Server-only service-role writer plus SQL schema"),
        ("Resend adapter", "Optional transactional email", "Defaults to in-app outbox unless live delivery is explicitly enabled"),
    ], [2160, 3600, 3600])

    add_h2(doc, "5.1 Data model")
    add_table(doc, ["Record", "Minimum fields", "Purpose"], [
        ("Recovery case", "ID, playbook, customer alias, at-risk amount, status, score", "Primary unit of recovery work"),
        ("Source event", "Provider event ID, type, timestamp, normalized payload", "Evidence for risk detection and webhook deduplication"),
        ("Policy decision", "Case ID, allowed flag, title, detail, next step", "Explains allow/defer/block behavior"),
        ("Intervention", "Action, channel, draft, approval, simulator/provider reference", "Records what was attempted"),
        ("Promise-to-pay", "Case ID, amount, due date, status", "Suppresses contact until resolution/expiry"),
        ("Recovery attribution", "Payment ID, case ID, confirmed amount, time", "Prevents double counting in metrics"),
        ("Audit event", "Time, type, title, detail, metadata", "Immutable case timeline and judge evidence"),
    ], [2160, 4080, 3120])

    add_h1(doc, "6. Interfaces and integrations")
    add_table(doc, ["Interface", "Contract", "Safety default"], [
        ("Razorpay webhook", "POST /api/webhooks/razorpay; validate HMAC signature and event ID", "Returns 503 unless a webhook secret is configured"),
        ("Razorpay Payment Links", "Server-side request using Test Mode credentials", "Returns a simulated link unless live outreach is explicitly enabled"),
        ("Resend", "Server-side transactional email request", "Saves a simulated outbox entry unless sender/key and explicit flag exist"),
        ("Supabase", "Server-side audit insert into recovery_audit_events", "Demo uses local in-memory data when no service role is configured"),
        ("Browser speech", "SpeechSynthesis preview of reviewed Hinglish script", "Does not place an outbound call or collect voice data"),
    ], [2160, 4320, 2880])

    add_h1(doc, "7. Non-functional requirements")
    add_table(doc, ["Category", "Requirement"], [
        ("Availability", "The public demo shall load with zero provider credentials and preserve the core judge flow locally."),
        ("Performance", "The synthetic dashboard shall render the 120-case batch and response details without a long-running background job."),
        ("Security", "Secrets shall remain server-side; no real customer data, card details, or production payment credentials shall be committed."),
        ("Privacy", "All default customer records, amounts, and outcomes shall be synthetic. UI shall label the workspace as demo-safe."),
        ("Accessibility", "The responsive interface shall provide readable contrast, semantic buttons, and keyboard-operable primary actions."),
        ("Observability", "The case audit trail shall explain every recommendation, action, stop rule, and outcome."),
        ("Deployability", "The repository shall support Vercel deployment from GitHub and local start with npm run dev."),
    ], [2160, 7200])

    add_h1(doc, "8. Test and acceptance plan")
    add_table(doc, ["Test layer", "Scenario", "Expected evidence"], [
        ("Unit", "Opt-out, contact cap, promise-to-pay, and settlement policy checks", "Policy engine allows or blocks with an exact visible reason"),
        ("Unit", "Safe payment recovery action", "Captured simulated payment produces one recovery attribution event"),
        ("Integration", "Webhook signature validation and missing configuration", "Invalid signature is rejected; unconfigured route cannot accept events"),
        ("End-to-end", "Every playbook selected and executed from the dashboard", "Audited outcome and batch metric update are visible"),
        ("Judge smoke test", "Open hosted app with no credentials", "Demo run starts, action works, blocked case is inspectable, metrics display"),
    ], [1800, 4080, 3480])
    add_h2(doc, "8.1 Acceptance criteria")
    add_body(doc, "The release is accepted when a judge can complete the following in under five minutes: start an isolated demo; inspect all seven recovery paths; execute at least one policy-approved action; inspect at least one blocked action; review audit evidence; and observe at-risk, recovered, and recovery-rate metrics.")

    add_h1(doc, "9. Deployment and handoff")
    add_table(doc, ["Deliverable", "Contents"], [
        ("Public GitHub repository", "Source, README, environment example, Supabase schema, tests, architecture, and video link"),
        ("Vercel deployment", "No-secret default demo; optional preview deployments from pull requests"),
        ("Five-minute pitch", "Problem, architecture, seven playbooks, safety block, audit evidence, and value metrics"),
        ("SRS", "This document, generated and visually checked as a Word deliverable"),
    ], [2700, 6660])
    add_h2(doc, "9.1 Future pilot boundary")
    add_body(doc, "A future merchant pilot would add tenant isolation, verified data-retention policy, provider-backed queues, genuine communication consent, legal/compliance review, real call/SMS integrations, and operational monitoring. None of those capabilities are represented as active production features in the hackathon release.")

    add_h1(doc, "10. Judge walkthrough")
    add_table(doc, ["Time", "Demonstration step"], [
        ("0:00-0:35", "State the revenue-recovery problem and open the 120-case overview."),
        ("0:35-1:10", "Explain event → AI recommendation → policy gate → adapter → reconciliation."),
        ("1:10-3:10", "Show payment degradation, checkout, subscriptions, B2B receivable, and mandate retry cases."),
        ("3:10-4:00", "Preview Hinglish voice and demonstrate promise-to-pay suppression."),
        ("4:00-4:40", "Show a policy-blocked action, audit timeline, and batch recovery metrics."),
        ("4:40-5:00", "Show GitHub, test command results, deployment path, and this requirements artifact."),
    ], [1800, 7560])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
